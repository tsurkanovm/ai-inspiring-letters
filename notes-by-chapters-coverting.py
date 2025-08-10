import os
import docx
import subprocess
import tempfile
from openai import OpenAI
import json

# Load config
with open("config.json") as f:
    config = json.load(f)

openai_client = OpenAI(api_key=config["openai_api_key"])

input_folder = config["notes_folder"]
output_folder = config["converted_notes_folder"]

os.makedirs(output_folder, exist_ok=True)


def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def convert_doc_to_docx(doc_path):
    tmp_dir = tempfile.mkdtemp()
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmp_dir, doc_path],
        check=True
    )
    converted_path = os.path.join(tmp_dir, os.path.splitext(os.path.basename(doc_path))[0] + ".docx")
    return converted_path


def create_docx(text, output_path):
    doc = docx.Document()
    for block in text.split("\n\n"):
        lines = block.strip().split("\n", 1)
        if len(lines) == 2:
            header, body = lines
            p = doc.add_paragraph()
            run = p.add_run(header.strip())
            run.bold = True
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
            run.font.size = docx.shared.Pt(14)
            p.style = 'Heading1'
            doc.add_paragraph(body.strip())
        else:
            doc.add_paragraph(block.strip())
    doc.save(output_path)


def transform_note_with_chatgpt(original_text):
    prompt = f"""
You are an assistant who helps reformat book notes into clean structured summaries.

Transform the following book notes with these rules:
1. Translate into English if needed (text can be partially or all in Russian).
2. REMOVE the main title/header.
3. Ignore any existing headers or markdown symbols.
4. DELETE all text between ** (they are old section titles).
5. Identify logical topic breaks and split into short sections.
6. Each section should have a NEW meaningful short header wrapped markdown (**) describing the idea, placed on its own line.
7. Only keep valuable content. Remove general or introductory parts like 'Introduction'.

The output must be:
- Clean readable plain text
- Each section starts with the NEW HEADER in uppercase, followed by the section's content.
- Headers must be separated from content by a new line.

Text to process:
\"\"\"
{original_text}
\"\"\"
"""
    response = openai_client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=4096,
        temperature=0.4
    )
    return response.choices[0].message.content.strip()


# Process all files
for filename in os.listdir(input_folder):
    if filename.lower().endswith((".docx", ".doc")):
        print(f"Processing: {filename}")
        input_path = os.path.join(input_folder, filename)

        if filename.lower().endswith(".doc"):
            input_path = convert_doc_to_docx(input_path)  # Convert first

        output_path = os.path.join(output_folder, os.path.splitext(filename)[0] + ".docx")

        original_text = read_docx(input_path)
        transformed_text = transform_note_with_chatgpt(original_text)
        create_docx(transformed_text, output_path)

        print(f"Saved: {output_path}")
